import os
import re
import io
import zipfile
import tempfile
import base64
import logging
import warnings
import asyncio
import shutil
from collections import deque

from telethon import TelegramClient, events, Button
from telethon.errors import MessageNotModifiedError
from ebooklib import epub
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from PIL import Image
from lxml import etree
from docx import Document

warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
logging.basicConfig(level=logging.INFO)

api_id = 24519852
api_hash = '2186f59fdf9c2ad4e7ddf0deb250ff0c'
client = TelegramClient('unified_bot', api_id, api_hash)

RESOLUTIONS = {
    'Удалить изображения': None,
    '64p': (64, 64),
    '144p': (256, 144),
    '360p': (640, 360),
    '480p': (854, 480),
    '720p': (1280, 720),
    '1080p': (1920, 1080)
}

user_files = {}
user_mode = {}
last_message_text = {}

# Система очереди
MAX_CONCURRENT_TASKS = 1
active_tasks = 0
task_queue = deque()
queue_lock = asyncio.Lock()

# Улучшенное регулярное выражение для поиска глав
CHAPTER_RE = re.compile(
    r"(?i)^\s*("
    r"Глава|chapter|часть|part|том|volume|книга|book|раздел|section|"
    r"пролог|prologue|эпилог|epilogue|предисловие|введение|заключение|"
    r"примечания|приложение|аннотация|annotation|описание"
    r")[\s\.:]*(\d+|[IVX]+|[А-Я])?", re.IGNORECASE | re.UNICODE
)

# Декоратор для управления очередью
async def queue_manager(func):
    """Декоратор для управления очередью обработки"""
    async def wrapper(*args, **kwargs):
        global active_tasks
        
        async with queue_lock:
            if active_tasks >= MAX_CONCURRENT_TASKS:
                future = asyncio.Future()
                task_queue.append(future)
                
                event = args[0]
                queue_position = len(task_queue)
                await event.respond(
                    f"⏳ Ваш запрос добавлен в очередь.\n"
                    f"Позиция в очереди: {queue_position}\n"
                    f"Ожидайте обработки..."
                )
                
                queue_lock.release()
                try:
                    await future
                finally:
                    await queue_lock.acquire()
            
            active_tasks += 1
        
        try:
            result = await func(*args, **kwargs)
            return result
        finally:
            async with queue_lock:
                active_tasks -= 1
                
                if task_queue:
                    next_task = task_queue.popleft()
                    next_task.set_result(True)
    
    return wrapper

# Функция для создания прогресс-бара
def create_progress_bar(current, total, width=20):
    if total == 0:
        return "▓" * width + " 100%"
    progress = int((current / total) * width)
    bar = "▓" * progress + "░" * (width - progress)
    percentage = int((current / total) * 100)
    return f"{bar} {percentage}%"

# Функция для проверки, нужно ли обновлять прогресс
def should_update_progress(current, total, last_updated_percent):
    if total == 0:
        return True
    current_percent = int((current / total) * 100)
    return current_percent - last_updated_percent >= 5

# Безопасное редактирование сообщения
async def safe_edit_message(message, new_text):
    """Редактирует сообщение только если текст изменился"""
    try:
        message_id = message.id
        if message_id not in last_message_text or last_message_text[message_id] != new_text:
            await message.edit(new_text)
            last_message_text[message_id] = new_text
    except MessageNotModifiedError:
        pass
    except Exception as e:
        logging.error(f"Error editing message: {e}")

# Команды
@client.on(events.NewMessage(pattern='/start'))
async def start(event):
    await event.respond(
        "Здравствуйте. Вас приветствует Адикия — бот для работы с документами.\n\n"
        "Возможности:\n"
        "• Сжатие или удаление изображений в .epub, .fb2, .docx\n"
        "• Извлечение глав из EPUB, FB2, DOCX и пересборка с оглавлением\n"
        "• Проверка и исправление ошибок в файлах\n\n"
        "⚠️ Бот работает на бесплатном сервере.\n"
        "Файлы обрабатываются строго по очереди.\n\n"
        "Используйте /help для получения списка команд."
    )

@client.on(events.NewMessage(pattern='/help'))
async def help_command(event):
    await event.respond(
        "Справка по командам:\n\n"
        "/start    - Начать работу с ботом\n"
        "/help     - Показать эту справку\n"
        "/compress - Сжать или удалить изображения\n"
        "/extract  - Извлечь главы и пересобрать с оглавлением\n"
        "/fix      - Проверка и исправление ошибок\n"
        "/cancel   - Отменить текущую операцию\n"
        "/status   - Проверить статус очереди\n"
    )

@client.on(events.NewMessage(pattern='/status'))
async def status_command(event):
    """Показывает текущий статус очереди"""
    status_text = "📊 Статус системы:\n\n"
    
    if active_tasks > 0:
        status_text += "🔄 Сейчас обрабатывается файл\n"
    else:
        status_text += "✅ Бот свободен\n"
    
    status_text += f"⏳ В очереди: {len(task_queue)} файл(ов)"
    
    await event.respond(status_text)

@client.on(events.NewMessage(pattern='/cancel'))
async def cancel(event):
    user_id = event.sender_id
    user_files.pop(user_id, None)
    user_mode.pop(user_id, None)
    await event.respond("Операция отменена.")

@client.on(events.NewMessage(pattern='/compress'))
async def compress_cmd(event):
    user_mode[event.sender_id] = 'compress'
    await event.respond("Пожалуйста, отправьте файл .epub, .fb2 или .docx для обработки изображений.")

@client.on(events.NewMessage(pattern='/extract'))
async def extract_cmd(event):
    user_mode[event.sender_id] = 'extract'
    await event.respond("Пожалуйста, отправьте файл .epub, .fb2 или .docx. Я извлеку главы и пересоберу его с оглавлением.")

@client.on(events.NewMessage(pattern='/fix'))
async def fix_cmd(event):
    user_mode[event.sender_id] = 'fix'
    await event.respond(
        "🔧 Функция проверки и починки файлов\n\n"
        "Отправьте файл .epub, .fb2 или .docx для диагностики и исправления ошибок."
    )

# Приём файлов с прогресс-баром
@client.on(events.NewMessage(incoming=True))
async def handle_file(event):
    if not event.file:
        return
    user_id = event.sender_id
    mode = user_mode.get(user_id)
    if not mode:
        return

    filename = event.file.name or ''
    ext = os.path.splitext(filename)[1].lower()
    file_size = event.file.size

    # Индикатор прогресса загрузки файла
    progress_msg = await event.respond("📥 Загрузка файла...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    
    file_data = io.BytesIO()
    last_percent = 0
    
    async def progress_callback(current, total):
        nonlocal last_percent
        if total > 0 and should_update_progress(current, total, last_percent):
            progress_bar = create_progress_bar(current, total)
            new_text = f"📥 Загрузка файла...\n{progress_bar}"
            await safe_edit_message(progress_msg, new_text)
            last_percent = int((current / total) * 100)
    
    await client.download_media(event.message, file=file_data, progress_callback=progress_callback)
    file_data.seek(0)

    await safe_edit_message(progress_msg, "📥 Загрузка завершена! Сохранение файла...")

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(file_data.read())
        tmp_path = tmp.name

    user_files[user_id] = (filename, tmp_path)
    await progress_msg.delete()
    last_message_text.pop(progress_msg.id, None)

    if mode == 'compress' and ext in ['.epub', '.fb2', '.docx']:
        buttons = [Button.inline(label, data=label.encode()) for label in RESOLUTIONS]
        await event.respond("Выберите способ обработки изображений:", buttons=buttons)

    elif mode == 'extract' and ext in ['.epub', '.fb2', '.docx']:
        await process_extract_with_queue(event, user_id, filename, tmp_path, ext)
        
    elif mode == 'fix' and ext in ['.epub', '.fb2', '.docx']:
        await process_fix_with_queue(event, user_id, filename, tmp_path, ext)

# Обработка извлечения с очередью
@queue_manager
async def process_extract_with_queue(event, user_id, filename, tmp_path, ext):
    """Обработка извлечения глав с учетом очереди"""
    await event.respond("✅ Файл принят в обработку. Начинаю...")
    try:
        base = os.path.splitext(filename)[0]
        output_path = None
        
        if ext == '.epub':
            chapters, images = await extract_chapters_from_epub_async(tmp_path, event)
            if not chapters:
                await event.respond("Главы не найдены в EPUB.")
                return
            output_path = os.path.join(tempfile.gettempdir(), f"{base}_converted.epub")
            build_progress = await event.respond("📚 Сборка EPUB...\n░░░░░░░░░░░░░░░░░░░░ 0%")
            await build_epub_async(base, chapters, images, output_path, build_progress)
            await build_progress.delete()
            last_message_text.pop(build_progress.id, None)
            await client.send_file(user_id, output_path, caption="✅ EPUB пересобран с оглавлением.")
            
        elif ext == '.fb2':
            chapters = await extract_chapters_from_fb2_async(tmp_path, event)
            if not chapters:
                await event.respond("Главы не найдены в FB2.")
                return
            output_path = os.path.join(tempfile.gettempdir(), f"{base}_converted.fb2")
            build_progress = await event.respond("📚 Сборка FB2...\n░░░░░░░░░░░░░░░░░░░░ 0%")
            await build_fb2_with_toc_async(base, chapters, output_path, build_progress)
            await build_progress.delete()
            last_message_text.pop(build_progress.id, None)
            await client.send_file(user_id, output_path, caption="✅ FB2 пересобран с оглавлением.")
            
        elif ext == '.docx':
            chapters = await extract_chapters_from_docx_async(tmp_path, event)
            if not chapters:
                await event.respond("Главы не найдены в DOCX.")
                return
            output_path = os.path.join(tempfile.gettempdir(), f"{base}_converted.docx")
            build_progress = await event.respond("📚 Сборка DOCX...\n░░░░░░░░░░░░░░░░░░░░ 0%")
            await build_docx_with_toc_async(base, chapters, output_path, build_progress)
            await build_progress.delete()
            last_message_text.pop(build_progress.id, None)
            await client.send_file(user_id, output_path, caption="✅ DOCX пересобран с оглавлением.")
        
        if output_path and os.path.exists(output_path):
            os.remove(output_path)
            
    except Exception as e:
        logging.error(f"Error processing file: {e}", exc_info=True)
        await event.respond(f"Ошибка: {e}")
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        user_mode.pop(user_id, None)
        user_files.pop(user_id, None)

# Обработка валидации с очередью
@queue_manager
async def process_fix_with_queue(event, user_id, filename, filepath, ext):
    """Обработка валидации и починки с учетом очереди"""
    await event.respond("🔧 Начинаю диагностику файла...")
    try:
        if ext == '.fb2':
            await process_fb2_fix(event, user_id, filename, filepath)
        elif ext == '.docx':
            await process_docx_fix(event, user_id, filename, filepath)
        elif ext == '.epub':
            await process_epub_fix(event, user_id, filename, filepath)
    finally:
        if os.path.exists(filepath):
            os.remove(filepath)
        user_files.pop(user_id, None)
        user_mode.pop(user_id, None)

# Inline-кнопки
@client.on(events.CallbackQuery)
async def handle_button(event):
    user_id = event.sender_id
    mode = user_mode.get(user_id)
    if not mode:
        return

    data = event.data.decode()
    filename, filepath = user_files.get(user_id, (None, None))
    if not filename or not os.path.exists(filepath):
        await event.answer("Файл не найден. Начните заново.", alert=True)
        return

    ext = os.path.splitext(filename)[1].lower()

    if mode == 'compress':
        resolution = RESOLUTIONS.get(data)
        await event.edit(f"⚙️ Подготовка к обработке файла {filename}...")
        await process_compression_with_queue(event, user_id, filename, filepath, resolution, ext)

# Обработка сжатия с очередью
@queue_manager
async def process_compression_with_queue(event, user_id, filename, filepath, resolution, ext):
    """Обработка сжатия изображений с учетом очереди"""
    await event.edit("✅ Файл принят в обработку. Начинаю...")
    
    if ext == '.fb2':
        await process_fb2(event, user_id, filename, filepath, resolution)
    elif ext == '.docx':
        await process_docx(event, user_id, filename, filepath, resolution)
    elif ext == '.epub':
        await process_epub_compression(event, user_id, filename, filepath, resolution)
    
    if os.path.exists(filepath):
        os.remove(filepath)
    user_files.pop(user_id, None)
    user_mode.pop(user_id, None)

# Обработка изображений FB2 с улучшенным прогресс-баром
async def process_fb2(event, user_id, filename, filepath, resolution):
    ns = {'fb2': 'http://www.gribuser.ru/xml/fictionbook/2.0'}
    tree = etree.parse(filepath)
    root = tree.getroot()
    binaries = root.xpath('//fb2:binary', namespaces=ns)

    changed = deleted = 0
    image_binaries = [b for b in binaries if 'image' in (b.get('content-type') or '')]
    total = len(image_binaries)
    
    if total == 0:
        await event.edit("Изображения не найдены в файле.")
        return

    progress_msg = await event.respond(f"🖼️ Обработка изображений FB2...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    last_percent = 0

    for current, binary in enumerate(image_binaries, 1):
        if should_update_progress(current, total, last_percent):
            progress_bar = create_progress_bar(current, total)
            action = "Удаление" if resolution is None else "Сжатие"
            new_text = f"🖼️ {action} изображений FB2...\n{progress_bar}"
            await safe_edit_message(progress_msg, new_text)
            last_percent = int((current / total) * 100)

        if resolution is None:
            binary_id = binary.get('id')
            if binary_id:
                for img_ref in root.xpath(f'//fb2:image[@l:href="#{binary_id}"]', 
                                        namespaces={'fb2': ns['fb2'], 'l': 'http://www.w3.org/1999/xlink'}):
                    img_ref.getparent().remove(img_ref)
            root.remove(binary)
            deleted += 1
        else:
            try:
                img = Image.open(io.BytesIO(base64.b64decode(binary.text))).convert('RGB')
                img.thumbnail(resolution, Image.Resampling.LANCZOS)
                buf = io.BytesIO()
                img.save(buf, format='JPEG', quality=30)
                binary.text = base64.b64encode(buf.getvalue()).decode()
                binary.set('content-type', 'image/jpeg')
                changed += 1
            except Exception:
                continue

    await safe_edit_message(progress_msg, "💾 Сохранение файла...")
    base, ext = os.path.splitext(filename)
    out_path = os.path.join(tempfile.gettempdir(), f"{base}_compressed{ext}")
    tree.write(out_path, encoding='utf-8', xml_declaration=True)
    
    await progress_msg.delete()
    last_message_text.pop(progress_msg.id, None)
    await client.send_file(user_id, out_path, caption=f"✅ Готово: сжато {changed}, удалено {deleted}")
    os.remove(filepath)
    os.remove(out_path)

# DOCX с улучшенным прогресс-баром
async def process_docx(event, user_id, filename, filepath, resolution):
    doc = Document(filepath)
    changed = deleted = 0
    total = len(doc.inline_shapes)
    
    if total == 0:
        await event.edit("Изображения не найдены в документе.")
        return

    progress_msg = await event.respond(f"🖼️ Обработка изображений DOCX...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    last_percent = 0

    if resolution is None:
        shapes_to_remove = []
        for current, shape in enumerate(doc.inline_shapes, 1):
            if should_update_progress(current, total, last_percent):
                progress_bar = create_progress_bar(current, total)
                new_text = f"🖼️ Удаление изображений DOCX...\n{progress_bar}"
                await safe_edit_message(progress_msg, new_text)
                last_percent = int((current / total) * 100)
            shapes_to_remove.append(shape)
        
        for shape in reversed(shapes_to_remove):
            try:
                shape._element.getparent().remove(shape._element)
                deleted += 1
            except Exception:
                continue
    else:
        for current, shape in enumerate(doc.inline_shapes, 1):
            if should_update_progress(current, total, last_percent):
                progress_bar = create_progress_bar(current, total)
                new_text = f"🖼️ Сжатие изображений DOCX...\n{progress_bar}"
                await safe_edit_message(progress_msg, new_text)
                last_percent = int((current / total) * 100)
            
            try:
                r_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                img_part = doc.part.related_parts[r_id]
                img = Image.open(io.BytesIO(img_part.blob)).convert('RGB')
                img.thumbnail(resolution, Image.Resampling.LANCZOS)
                buf = io.BytesIO()
                img.save(buf, format='JPEG', quality=30)
                img_part._blob = buf.getvalue()
                changed += 1
            except Exception:
                continue

    await safe_edit_message(progress_msg, "💾 Сохранение документа...")
    base, ext = os.path.splitext(filename)
    out_path = os.path.join(tempfile.gettempdir(), f"{base}_compressed{ext}")
    doc.save(out_path)
    
    await progress_msg.delete()
    last_message_text.pop(progress_msg.id, None)
    await client.send_file(user_id, out_path, caption=f"✅ Готово: сжато {changed}, удалено {deleted}")
    os.remove(filepath)
    os.remove(out_path)

# EPUB с улучшенным прогресс-баром
async def process_epub_compression(event, user_id, filename, filepath, resolution):
    book = epub.read_epub(filepath)
    changed = deleted = 0
    images = [item for item in list(book.get_items()) if item.media_type and item.media_type.startswith("image/")]
    total = len(images)
    
    if total == 0:
        await event.edit("Изображения не найдены в EPUB.")
        return

    progress_msg = await event.respond(f"🖼️ Обработка изображений EPUB...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    last_percent = 0

    for current, item in enumerate(images, 1):
        if should_update_progress(current, total, last_percent):
            progress_bar = create_progress_bar(current, total)
            action = "Удаление" if resolution is None else "Сжатие"
            new_text = f"🖼️ {action} изображений EPUB...\n{progress_bar}"
            await safe_edit_message(progress_msg, new_text)
            last_percent = int((current / total) * 100)

        if resolution is None:
            book.items.remove(item)
            deleted += 1
        else:
            try:
                img = Image.open(io.BytesIO(item.content)).convert("RGB")
                img.thumbnail(resolution, Image.Resampling.LANCZOS)
                buf = io.BytesIO()
                img.save(buf, format="JPEG", quality=30)
                item.content = buf.getvalue()
                item.media_type = "image/jpeg"
                changed += 1
            except Exception:
                continue

    await safe_edit_message(progress_msg, "💾 Сохранение EPUB...")
    base, ext = os.path.splitext(filename)
    out_path = os.path.join(tempfile.gettempdir(), f"{base}_compressed{ext}")
    epub.write_epub(out_path, book)
    
    await progress_msg.delete()
    last_message_text.pop(progress_msg.id, None)
    await client.send_file(user_id, out_path, caption=f"✅ Готово: сжато {changed}, удалено {deleted}")
    os.remove(filepath)
    os.remove(out_path)

# EPUB: главы с асинхронным прогресс-баром
async def extract_chapters_from_epub_async(epub_path, event):
    temp_dir = tempfile.mkdtemp()
    html_blocks = []
    images = {}

    progress_msg = await event.respond("📖 Извлечение файлов из EPUB...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    
    try:
        with zipfile.ZipFile(epub_path, 'r') as zf:
            zf.extractall(temp_dir)

        await safe_edit_message(progress_msg, "📖 Анализ содержимого...\n▓▓▓▓▓░░░░░░░░░░░░░░░ 25%")
        
        all_files = []
        for root, _, files in os.walk(temp_dir):
            for file in files:
                all_files.append(os.path.join(root, file))
        
        processed = 0
        total_files = len(all_files)
        
        for file_path in all_files:
            file = os.path.basename(file_path)
            if file.lower().endswith((".xhtml", ".html", ".htm")):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        html_blocks.append(BeautifulSoup(f, "lxml"))
                except Exception as e:
                    logging.warning(f"Could not parse HTML file {file}: {e}")
            elif file.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".svg", ".webp")):
                images[file] = file_path
            
            processed += 1
            if processed % 10 == 0:
                progress = int((processed / total_files) * 25) + 25
                progress_bar = "▓" * (progress // 5) + "░" * (20 - progress // 5)
                new_text = f"📖 Анализ файлов...\n{progress_bar} {progress}%"
                await safe_edit_message(progress_msg, new_text)

        await safe_edit_message(progress_msg, "🔍 Поиск глав...\n▓▓▓▓▓▓▓▓▓▓░░░░░░░░░░ 50%")
        
        chapters = []
        title, content, num = None, "", 0
        
        processed_blocks = 0
        total_blocks = len(html_blocks)
        
        for soup in html_blocks:
            if not soup.body:
                continue
            for elem in soup.body.find_all(recursive=False):
                text = elem.get_text(strip=True)
                match = CHAPTER_RE.match(text or "")
                if match:
                    if title:
                        chapters.append((num, title, content.strip()))
                    title = text
                    num_match = re.search(r'\d+', title)
                    num = int(num_match.group()) if num_match else len(chapters)
                    content = f"<h1>{title}</h1>"
                else:
                    content += str(elem)
            
            processed_blocks += 1
            if processed_blocks % 5 == 0:
                progress = 50 + int((processed_blocks / total_blocks) * 25)
                progress_bar = "▓" * (progress // 5) + "░" * (20 - progress // 5)
                new_text = f"🔍 Поиск глав...\n{progress_bar} {progress}%"
                await safe_edit_message(progress_msg, new_text)
        
        if title:
            chapters.append((num, title, content.strip()))

        await safe_edit_message(progress_msg, "🔄 Обработка результатов...\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓░░░░░ 75%")
        
        seen, result = set(), []
        for n, t, c in sorted(chapters, key=lambda x: x[0]):
            if t not in seen and len(c.strip()) > len(f"<h1>{t}</h1>"):
                result.append((n, t, c))
                seen.add(t)
        
        await safe_edit_message(progress_msg, f"✅ Найдено глав: {len(result)}\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓ 100%")
        await asyncio.sleep(1)
        await progress_msg.delete()
        last_message_text.pop(progress_msg.id, None)
        
        return result, images
        
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

async def build_epub_async(title, chapters, image_paths, output_path, progress_msg):
    book = epub.EpubBook()
    book.set_identifier("converted")
    book.set_title(title)
    book.set_language("ru")
    book.add_author("Adikia Bot")

    spine = ['nav']
    toc = []
    
    total_steps = len(image_paths) + len(chapters) + 3
    current_step = 0

    for fname, path in image_paths.items():
        current_step += 1
        progress_bar = create_progress_bar(current_step, total_steps)
        new_text = f"📚 Добавление изображений...\n{progress_bar}"
        await safe_edit_message(progress_msg, new_text)
        
        try:
            ext = os.path.splitext(fname)[1][1:].lower()
            mime = f"image/{'jpeg' if ext in ['jpg', 'jpeg'] else ext}"
            with open(path, 'rb') as f:
                book.add_item(epub.EpubItem(uid=fname, file_name=f"images/{fname}", media_type=mime, content=f.read()))
        except Exception as e:
            logging.warning(f"Could not add image {fname}: {e}")

    for i, (num, chapter_title, html_body) in enumerate(chapters, 1):
        current_step += 1
        progress_bar = create_progress_bar(current_step, total_steps)
        new_text = f"📚 Добавление глав...\n{progress_bar}"
        await safe_edit_message(progress_msg, new_text)
        
        html = epub.EpubHtml(title=chapter_title, file_name=f"chap_{i}.xhtml", lang='ru')
        html.content = html_body
        book.add_item(html)
        spine.append(html)
        toc.append(epub.Link(html.file_name, chapter_title, f"chap_{i}"))

    current_step += 1
    progress_bar = create_progress_bar(current_step, total_steps)
    await safe_edit_message(progress_msg, f"📚 Создание оглавления...\n{progress_bar}")
    
    book.spine = spine
    book.toc = toc
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    
    current_step += 1
    progress_bar = create_progress_bar(current_step, total_steps)
    await safe_edit_message(progress_msg, f"📚 Сохранение EPUB...\n{progress_bar}")
    
    epub.write_epub(output_path, book)
    
    current_step += 1
    progress_bar = create_progress_bar(current_step, total_steps)
    await safe_edit_message(progress_msg, f"📚 Завершение...\n{progress_bar}")

# УЛУЧШЕННАЯ функция для извлечения глав из FB2
async def extract_chapters_from_fb2_async(fb2_path, event):
    """Улучшенное извлечение глав из FB2 файла"""
    progress_msg = await event.respond("📖 Анализ структуры FB2...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    
    try:
        ns = {'fb2': 'http://www.gribuser.ru/xml/fictionbook/2.0'}
        tree = etree.parse(fb2_path)
        root = tree.getroot()
        
        await safe_edit_message(progress_msg, "📖 Поиск глав в FB2...\n▓▓▓▓▓░░░░░░░░░░░░░░░ 25%")
        
        chapters = []
        
        # Метод 1: Ищем секции с заголовками
        sections = root.xpath('//fb2:section[fb2:title]', namespaces=ns)
        
        # Метод 2: Если секций мало, ищем все параграфы с паттернами глав
        if len(sections) < 3:
            await safe_edit_message(progress_msg, "🔍 Глубокий поиск глав...\n▓▓▓▓▓▓▓▓░░░░░░░░░░░░ 40%")
            
            # Собираем все текстовые элементы
            all_paragraphs = root.xpath('//fb2:p', namespaces=ns)
            current_chapter = None
            current_content = []
            chapter_num = 0
            
            for para in all_paragraphs:
                text = ''.join(para.itertext()).strip()
                
                # Проверяем, является ли параграф заголовком главы
                if text and CHAPTER_RE.match(text):
                    # Сохраняем предыдущую главу
                    if current_chapter and current_content:
                        # Создаем XML структуру для главы
                        section_xml = f'<section xmlns="{ns["fb2"]}">'
                        section_xml += f'<title><p>{current_chapter}</p></title>'
                        section_xml += ''.join(current_content)
                        section_xml += '</section>'
                        
                        chapters.append((chapter_num, current_chapter, section_xml))
                    
                    # Начинаем новую главу
                    current_chapter = text
                    num_match = re.search(r'\d+', text)
                    chapter_num = int(num_match.group()) if num_match else chapter_num + 1
                    current_content = []
                elif current_chapter:
                    # Добавляем параграф к текущей главе
                    para_str = etree.tostring(para, encoding='unicode', pretty_print=True)
                    current_content.append(para_str)
            
            # Добавляем последнюю главу
            if current_chapter and current_content:
                section_xml = f'<section xmlns="{ns["fb2"]}">'
                section_xml += f'<title><p>{current_chapter}</p></title>'
                section_xml += ''.join(current_content)
                section_xml += '</section>'
                
                chapters.append((chapter_num, current_chapter, section_xml))
        
        else:
            # Обрабатываем найденные секции
            total = len(sections)
            
            for idx, section in enumerate(sections):
                if idx % 5 == 0:
                    progress = 25 + int((idx / total) * 50)
                    progress_bar = "▓" * (progress // 5) + "░" * (20 - progress // 5)
                    await safe_edit_message(progress_msg, f"📖 Обработка секций...\n{progress_bar} {progress}%")
                
                # Получаем заголовок
                title_elem = section.find('.//fb2:title', namespaces=ns)
                if title_elem is not None:
                    # Извлекаем текст заголовка, учитывая вложенные элементы
                    title_text = ' '.join(title_elem.itertext()).strip()
                    
                    # Проверяем, не является ли это просто номером
                    if title_text and (len(title_text) > 2 or not title_text.isdigit()):
                        # Извлекаем содержимое секции
                        content = etree.tostring(section, encoding='unicode', pretty_print=True)
                        
                        # Извлекаем номер главы если есть
                        num_match = re.search(r'\d+', title_text)
                        num = int(num_match.group()) if num_match else idx
                        
                        chapters.append((num, title_text, content))
        
        # Проверяем вложенные секции
        if len(chapters) < 3:
            await safe_edit_message(progress_msg, "🔍 Поиск вложенных глав...\n▓▓▓▓▓▓▓▓▓▓▓▓▓░░░░░░░ 65%")
            
            # Ищем секции внутри body
            body_sections = root.xpath('//fb2:body/fb2:section', namespaces=ns)
            for section in body_sections:
                # Ищем вложенные секции с заголовками
                nested_sections = section.xpath('.//fb2:section[fb2:title]', namespaces=ns)
                
                for nested in nested_sections:
                    title_elem = nested.find('.//fb2:title', namespaces=ns)
                    if title_elem is not None:
                        title_text = ' '.join(title_elem.itertext()).strip()
                        
                        if title_text and CHAPTER_RE.match(title_text):
                            content = etree.tostring(nested, encoding='unicode', pretty_print=True)
                            num_match = re.search(r'\d+', title_text)
                            num = int(num_match.group()) if num_match else len(chapters)
                            
                            # Проверяем, что эта глава еще не добавлена
                            if not any(ch[1] == title_text for ch in chapters):
                                chapters.append((num, title_text, content))
        
        # Сортируем главы по номерам
        chapters.sort(key=lambda x: x[0])
        
        # Убираем дубликаты
        unique_chapters = []
        seen_titles = set()
        for num, title, content in chapters:
            if title not in seen_titles:
                unique_chapters.append((num, title, content))
                seen_titles.add(title)
        
        await safe_edit_message(progress_msg, f"✅ Найдено глав: {len(unique_chapters)}\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓ 100%")
        await asyncio.sleep(1)
        await progress_msg.delete()
        last_message_text.pop(progress_msg.id, None)
        
        return unique_chapters
        
    except Exception as e:
        logging.error(f"Error parsing FB2: {e}")
        await safe_edit_message(progress_msg, f"❌ Ошибка при обработке FB2: {e}")
        await asyncio.sleep(3)
        await progress_msg.delete()
        return []

# Функция для создания FB2 с оглавлением
async def build_fb2_with_toc_async(title, chapters, output_path, progress_msg):
    """Создает FB2 файл с правильным оглавлением"""
    ns = {'fb2': 'http://www.gribuser.ru/xml/fictionbook/2.0'}
    
    root = etree.Element('{http://www.gribuser.ru/xml/fictionbook/2.0}FictionBook', 
                        nsmap={None: ns['fb2']})
    root.set('{http://www.w3.org/2001/XMLSchema-instance}schemaLocation', 
             'http://www.gribuser.ru/xml/fictionbook/2.0 http://www.gribuser.ru/xml/fictionbook/2.0/FictionBook2.xsd')
    
    description = etree.SubElement(root, 'description')
    title_info = etree.SubElement(description, 'title-info')
    etree.SubElement(title_info, 'book-title').text = title
    etree.SubElement(title_info, 'lang').text = 'ru'
    
    body = etree.SubElement(root, 'body')
    
    total = len(chapters)
    
    for idx, (num, chapter_title, content) in enumerate(chapters):
        progress = int((idx / total) * 100)
        progress_bar = create_progress_bar(idx, total)
        await safe_edit_message(progress_msg, f"📚 Сборка FB2...\n{progress_bar}")
        
        try:
            content = content.replace('<?xml version="1.0" encoding="utf-8"?>', '')
            section = etree.fromstring(content)
            for elem in section.iter():
                elem.tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            body.append(section)
        except Exception as e:
            logging.error(f"Could not parse chapter content: {e}")
            section = etree.SubElement(body, 'section')
            section_title = etree.SubElement(section, 'title')
            p = etree.SubElement(section_title, 'p')
            p.text = chapter_title
    
    await safe_edit_message(progress_msg, "💾 Сохранение FB2...\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓ 100%")
    tree = etree.ElementTree(root)
    tree.write(output_path, encoding='utf-8', xml_declaration=True, pretty_print=True)

# Функция для извлечения глав из DOCX
async def extract_chapters_from_docx_async(docx_path, event):
    """Извлекает главы из DOCX файла по заголовкам"""
    progress_msg = await event.respond("📖 Анализ структуры DOCX...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    
    try:
        doc = Document(docx_path)
        chapters = []
        current_chapter = None
        current_content = []
        
        total = len(doc.paragraphs)
        
        for idx, para in enumerate(doc.paragraphs):
            if idx % 20 == 0:
                progress = int((idx / total) * 75)
                progress_bar = create_progress_bar(idx, total)
                await safe_edit_message(progress_msg, f"📖 Анализ параграфов...\n{progress_bar}")
            
            text = para.text.strip()
            if text and (para.style.name.startswith('Heading') or CHAPTER_RE.match(text)):
                if current_chapter and current_content:
                    num_match = re.search(r'\d+', current_chapter)
                    num = int(num_match.group()) if num_match else len(chapters)
                    chapters.append((num, current_chapter, '\n'.join(current_content)))
                
                current_chapter = text
                current_content = []
            elif current_chapter and text:
                current_content.append(text)
        
        if current_chapter and current_content:
            num_match = re.search(r'\d+', current_chapter)
            num = int(num_match.group()) if num_match else len(chapters)
            chapters.append((num, current_chapter, '\n'.join(current_content)))
        
        await safe_edit_message(progress_msg, f"✅ Найдено глав: {len(chapters)}\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓ 100%")
        await asyncio.sleep(1)
        await progress_msg.delete()
        last_message_text.pop(progress_msg.id, None)
        
        return chapters
        
    except Exception as e:
        logging.error(f"Error parsing DOCX: {e}")
        await safe_edit_message(progress_msg, f"❌ Ошибка при обработке DOCX: {e}")
        await asyncio.sleep(3)
        await progress_msg.delete()
        return []

# Функция для создания DOCX с оглавлением
async def build_docx_with_toc_async(title, chapters, output_path, progress_msg):
    """Создает DOCX файл с оглавлением"""
    doc = Document()
    
    doc.add_heading(title, 0)
    doc.add_heading('Оглавление', 1)
    
    total_steps = len(chapters) * 2 + 2
    current_step = 0
    
    for num, chapter_title, _ in sorted(chapters, key=lambda x: x[0]):
        current_step += 1
        progress_bar = create_progress_bar(current_step, total_steps)
        await safe_edit_message(progress_msg, f"📚 Создание оглавления...\n{progress_bar}")
        
        doc.add_paragraph(f"{num + 1}. {chapter_title}", style='List Number')
    
    doc.add_page_break()
    
    for num, chapter_title, content in sorted(chapters, key=lambda x: x[0]):
        current_step += 1
        progress_bar = create_progress_bar(current_step, total_steps)
        await safe_edit_message(progress_msg, f"📚 Добавление глав...\n{progress_bar}")
        
        doc.add_heading(chapter_title, 1)
        
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph)
    
    current_step += 1
    progress_bar = create_progress_bar(current_step, total_steps)
    await safe_edit_message(progress_msg, f"💾 Сохранение DOCX...\n{progress_bar}")
    
    doc.save(output_path)

# Функции валидации и починки

async def process_fb2_fix(event, user_id, filename, filepath):
    """Валидация и починка FB2"""
    report = "📋 Отчет о проверке FB2:\n\n"
    fixed_issues = []
    
    progress_msg = await event.respond("🔧 Анализ структуры FB2...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    
    try:
        ns = {'fb2': 'http://www.gribuser.ru/xml/fictionbook/2.0'}
        tree = etree.parse(filepath)
        root = tree.getroot()
        
        await safe_edit_message(progress_msg, "🔧 Проверка структуры...\n▓▓▓▓░░░░░░░░░░░░░░░░ 20%")
        
        if not root.find('.//fb2:description', namespaces=ns):
            report += "❌ Отсутствует элемент description\n"
            description = etree.SubElement(root, '{%s}description' % ns['fb2'])
            title_info = etree.SubElement(description, '{%s}title-info' % ns['fb2'])
            etree.SubElement(title_info, '{%s}book-title' % ns['fb2']).text = "Без названия"
            etree.SubElement(title_info, '{%s}lang' % ns['fb2']).text = "ru"
            fixed_issues.append("Добавлен элемент description")
        
        await safe_edit_message(progress_msg, "🔧 Проверка изображений...\n▓▓▓▓▓▓▓▓░░░░░░░░░░░░ 40%")
        
        image_refs = root.xpath('//fb2:image', namespaces=ns)
        binaries = {b.get('id'): b for b in root.xpath('//fb2:binary', namespaces=ns)}
        
        broken_images = 0
        for img in image_refs:
            href = img.get('{http://www.w3.org/1999/xlink}href')
            if href and href.startswith('#'):
                binary_id = href[1:]
                if binary_id not in binaries:
                    img.getparent().remove(img)
                    broken_images += 1
        
        if broken_images > 0:
            report += f"🖼️ Удалено битых ссылок на изображения: {broken_images}\n"
            fixed_issues.append(f"Удалено {broken_images} битых изображений")
        
        await safe_edit_message(progress_msg, "🔧 Проверка пустых элементов...\n▓▓▓▓▓▓▓▓▓▓▓▓░░░░░░░ 60%")
        
        empty_elements = 0
        for elem in root.xpath('.//*', namespaces=ns):
            if elem.tag.endswith('p') and not elem.text and not elem.tail and len(elem) == 0:
                elem.getparent().remove(elem)
                empty_elements += 1
        
        if empty_elements > 0:
            report += f"📄 Удалено пустых параграфов: {empty_elements}\n"
            fixed_issues.append(f"Удалено {empty_elements} пустых элементов")
        
        await safe_edit_message(progress_msg, "🔧 Проверка кодировки...\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓░░░░ 80%")
        
        for elem in root.xpath('.//*', namespaces=ns):
            if elem.text:
                elem.text = elem.text.replace('РІ', 'в').replace('вЂ"', '—').replace('вЂ™', ''')
            if elem.tail:
                elem.tail = elem.tail.replace('РІ', 'в').replace('вЂ"', '—').replace('вЂ™', ''')
        
        await safe_edit_message(progress_msg, "💾 Сохранение исправленного файла...")
        
        if fixed_issues:
            base, ext = os.path.splitext(filename)
            out_path = os.path.join(tempfile.gettempdir(), f"{base}_fixed{ext}")
            tree.write(out_path, encoding='utf-8', xml_declaration=True, pretty_print=True)
            
            report += f"\n✅ Исправлено проблем: {len(fixed_issues)}\n"
            report += "\n".join(f"• {issue}" for issue in fixed_issues)
            
            await progress_msg.delete()
            await client.send_file(user_id, out_path, caption=report)
            os.remove(out_path)
        else:
            report += "\n✅ Проблем не обнаружено!"
            await progress_msg.delete()
            await event.respond(report)
            
    except etree.XMLSyntaxError as e:
        await progress_msg.delete()
        await event.respond(
            f"❌ Критическая ошибка XML:\n{str(e)}\n\n"
            "Файл серьезно поврежден и требует ручного исправления."
        )
    except Exception as e:
        await progress_msg.delete()
        await event.respond(f"❌ Ошибка при проверке: {e}")

async def process_docx_fix(event, user_id, filename, filepath):
    """Валидация и починка DOCX"""
    report = "📋 Отчет о проверке DOCX:\n\n"
    fixed_issues = []
    
    progress_msg = await event.respond("🔧 Анализ документа DOCX...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    
    try:
        doc = Document(filepath)
        
        await safe_edit_message(progress_msg, "🔧 Проверка пустых параграфов...\n▓▓▓▓▓░░░░░░░░░░░░░░░ 25%")
        
        empty_paras = 0
        paras_to_remove = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                has_images = any(run._element.xpath('.//w:drawing') for run in para.runs)
                if not has_images:
                    paras_to_remove.append(para)
                    empty_paras += 1
        
        for para in paras_to_remove:
            p = para._element
            p.getparent().remove(p)
        
        if empty_paras > 0:
            report += f"📄 Удалено пустых параграфов: {empty_paras}\n"
            fixed_issues.append(f"Удалено {empty_paras} пустых параграфов")
        
        await safe_edit_message(progress_msg, "🔧 Проверка стилей...\n▓▓▓▓▓▓▓▓▓▓░░░░░░░░░░ 50%")
        
        fixed_styles = 0
        for para in doc.paragraphs:
            if para.style is None or para.style.name == 'Normal':
                if CHAPTER_RE.match(para.text):
                    para.style = 'Heading 1'
                    fixed_styles += 1
        
        if fixed_styles > 0:
            report += f"🎨 Исправлено стилей заголовков: {fixed_styles}\n"
            fixed_issues.append(f"Исправлено {fixed_styles} стилей")
        
        await safe_edit_message(progress_msg, "🔧 Проверка изображений...\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓░░░░░ 75%")
        
        broken_images = 0
        for shape in doc.inline_shapes:
            try:
                _ = shape._inline.graphic.graphicData
            except:
                broken_images += 1
        
        if broken_images > 0:
            report += f"🖼️ Обнаружено битых изображений: {broken_images}\n"
        
        await safe_edit_message(progress_msg, "🔧 Проверка списков...\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓ 100%")
        
        total_paras = len(doc.paragraphs)
        total_tables = len(doc.tables)
        total_sections = len(doc.sections)
        
        report += f"\n📊 Статистика документа:\n"
        report += f"• Параграфов: {total_paras}\n"
        report += f"• Таблиц: {total_tables}\n"
        report += f"• Секций: {total_sections}\n"
        
        if fixed_issues:
            base, ext = os.path.splitext(filename)
            out_path = os.path.join(tempfile.gettempdir(), f"{base}_fixed{ext}")
            doc.save(out_path)
            
            report += f"\n✅ Исправлено проблем: {len(fixed_issues)}\n"
            report += "\n".join(f"• {issue}" for issue in fixed_issues)
            
            await progress_msg.delete()
            await client.send_file(user_id, out_path, caption=report)
            os.remove(out_path)
        else:
            report += "\n✅ Проблем не обнаружено!"
            await progress_msg.delete()
            await event.respond(report)
            
    except Exception as e:
        await progress_msg.delete()
        await event.respond(f"❌ Ошибка при проверке: {e}")

async def process_epub_fix(event, user_id, filename, filepath):
    """Валидация и починка EPUB"""
    report = "📋 Отчет о проверке EPUB:\n\n"
    fixed_issues = []
    
    progress_msg = await event.respond("🔧 Анализ структуры EPUB...\n░░░░░░░░░░░░░░░░░░░░ 0%")
    
    try:
        book = epub.read_epub(filepath)
        
        await safe_edit_message(progress_msg, "🔧 Проверка метаданных...\n▓▓▓▓░░░░░░░░░░░░░░░░ 20%")
        
        if not book.get_metadata('DC', 'title'):
            book.set_title('Без названия')
            fixed_issues.append("Добавлено название книги")
            
        if not book.get_metadata('DC', 'language'):
            book.set_language('ru')
            fixed_issues.append("Добавлен язык")
            
        if not book.get_metadata('DC', 'identifier'):
            book.set_identifier(f'urn:uuid:{os.urandom(16).hex()}')
            fixed_issues.append("Добавлен идентификатор")
        
        await safe_edit_message(progress_msg, "🔧 Проверка изображений...\n▓▓▓▓▓▓▓▓░░░░░░░░░░░░ 40%")
        
        images = list(book.get_items_of_type(ebooklib.ITEM_IMAGE))
        broken_images = []
        
        for img in images:
            try:
                Image.open(io.BytesIO(img.content))
            except:
                broken_images.append(img)
        
        if broken_images:
            for img in broken_images:
                book.items.remove(img)
            report += f"🖼️ Удалено битых изображений: {len(broken_images)}\n"
            fixed_issues.append(f"Удалено {len(broken_images)} битых изображений")
        
        await safe_edit_message(progress_msg, "🔧 Проверка HTML контента...\n▓▓▓▓▓▓▓▓▓▓▓▓░░░░░░░ 60%")
        
        documents = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
        fixed_html = 0
        
        for item in documents:
            try:
                content = item.get_content().decode('utf-8')
                soup = BeautifulSoup(content, 'html.parser')
                
                for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    if not tag.get_text(strip=True):
                        tag.decompose()
                        fixed_html += 1
                
                for img in soup.find_all('img'):
                    src = img.get('src')
                    if src:
                        img_name = os.path.basename(src)
                        if not any(i.file_name.endswith(img_name) for i in images):
                            img.decompose()
                            fixed_html += 1
                
                item.set_content(str(soup).encode('utf-8'))
                
            except Exception as e:
                report += f"⚠️ Ошибка в файле {item.file_name}: {str(e)}\n"
        
        if fixed_html > 0:
            report += f"📄 Исправлено HTML элементов: {fixed_html}\n"
            fixed_issues.append(f"Исправлено {fixed_html} HTML элементов")
        
        await safe_edit_message(progress_msg, "🔧 Проверка навигации...\n▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓░░░░ 80%")
        
        spine_items = []
        for item_id, linear in book.spine:
            try:
                item = book.get_item_with_id(item_id)
                if item:
                    spine_items.append((item_id, linear))
            except:
                fixed_issues.append(f"Удалена битая ссылка из spine: {item_id}")
        
        book.spine = spine_items
        
        report += f"\n📊 Статистика EPUB:\n"
        report += f"• HTML файлов: {len(documents)}\n"
        report += f"• Изображений: {len(images) - len(broken_images)}\n"
        report += f"• Элементов в spine: {len(spine_items)}\n"
        
        if fixed_issues:
            await safe_edit_message(progress_msg, "💾 Сохранение исправленного EPUB...")
            base, ext = os.path.splitext(filename)
            out_path = os.path.join(tempfile.gettempdir(), f"{base}_fixed{ext}")
            epub.write_epub(out_path, book)
            
            report += f"\n✅ Исправлено проблем: {len(fixed_issues)}\n"
            report += "\n".join(f"• {issue}" for issue in fixed_issues)
            
            await progress_msg.delete()
            await client.send_file(user_id, out_path, caption=report)
            os.remove(out_path)
        else:
            report += "\n✅ Проблем не обнаружено!"
            await progress_msg.delete()
            await event.respond(report)
            
    except Exception as e:
        await progress_msg.delete()
        await event.respond(f"❌ Ошибка при проверке: {e}")

# Запуск
client.start()
print("Бот запущен.")
print("Обработка файлов: строго по очереди")
client.run_until_disconnected()
